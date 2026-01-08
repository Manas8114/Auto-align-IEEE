"""
IEEE-Compliant DOCX Generator.

This module generates properly formatted Word documents following
IEEE conference paper guidelines using python-docx.
"""

import io
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models.schemas import StructuredPaper, Section, Figure


class IEEEDocxGenerator:
    """Generates IEEE-formatted Word documents."""
    
    # IEEE Formatting Constants
    PAGE_WIDTH = Inches(8.27)  # A4 width
    PAGE_HEIGHT = Inches(11.69)  # A4 height
    MARGIN_TOP = Inches(0.75)
    MARGIN_BOTTOM = Inches(1.0)
    MARGIN_LEFT = Inches(0.625)
    MARGIN_RIGHT = Inches(0.625)
    COLUMN_GAP = Inches(0.25)
    
    FONT_TITLE = "Times New Roman"
    FONT_SIZE_TITLE = Pt(24)
    FONT_SIZE_AUTHOR = Pt(11)
    FONT_SIZE_ABSTRACT = Pt(9)
    FONT_SIZE_BODY = Pt(10)
    FONT_SIZE_HEADING = Pt(10)
    FONT_SIZE_CAPTION = Pt(8)
    
    def __init__(self):
        """Initialize the document generator."""
        self.doc = None
        self.figure_counter = 0
        self.table_counter = 0
        
    def _setup_document(self):
        """Set up document with IEEE formatting."""
        self.doc = Document()
        self.figure_counter = 0
        self.table_counter = 0
        
        # Set up page size (A4)
        section = self.doc.sections[0]
        section.page_width = self.PAGE_WIDTH
        section.page_height = self.PAGE_HEIGHT
        section.top_margin = self.MARGIN_TOP
        section.bottom_margin = self.MARGIN_BOTTOM
        section.left_margin = self.MARGIN_LEFT
        section.right_margin = self.MARGIN_RIGHT
        
        # Create custom styles
        self._create_styles()
    
    def _create_styles(self):
        """Create IEEE-specific paragraph styles."""
        styles = self.doc.styles
        
        # Title Style
        if "IEEE Title" not in [s.name for s in styles]:
            title_style = styles.add_style("IEEE Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.FONT_TITLE
            title_style.font.size = self.FONT_SIZE_TITLE
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Author Style
        if "IEEE Author" not in [s.name for s in styles]:
            author_style = styles.add_style("IEEE Author", WD_STYLE_TYPE.PARAGRAPH)
            author_style.font.name = self.FONT_TITLE
            author_style.font.size = self.FONT_SIZE_AUTHOR
            author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_style.paragraph_format.space_after = Pt(12)
        
        # Abstract Label Style
        if "IEEE Abstract Label" not in [s.name for s in styles]:
            abs_label = styles.add_style("IEEE Abstract Label", WD_STYLE_TYPE.PARAGRAPH)
            abs_label.font.name = self.FONT_TITLE
            abs_label.font.size = self.FONT_SIZE_ABSTRACT
            abs_label.font.bold = True
            abs_label.font.italic = True
            abs_label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abs_label.paragraph_format.space_after = Pt(6)
        
        # Abstract Text Style
        if "IEEE Abstract" not in [s.name for s in styles]:
            abs_style = styles.add_style("IEEE Abstract", WD_STYLE_TYPE.PARAGRAPH)
            abs_style.font.name = self.FONT_TITLE
            abs_style.font.size = self.FONT_SIZE_ABSTRACT
            abs_style.font.italic = True
            abs_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abs_style.paragraph_format.space_after = Pt(12)
        
        # Section Heading Style (Small Caps simulated with uppercase)
        if "IEEE Heading" not in [s.name for s in styles]:
            heading_style = styles.add_style("IEEE Heading", WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = self.FONT_TITLE
            heading_style.font.size = self.FONT_SIZE_HEADING
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Body Text Style
        if "IEEE Body" not in [s.name for s in styles]:
            body_style = styles.add_style("IEEE Body", WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = self.FONT_TITLE
            body_style.font.size = self.FONT_SIZE_BODY
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.first_line_indent = Inches(0.25)
            body_style.paragraph_format.space_after = Pt(0)
            body_style.paragraph_format.line_spacing = 1.0
        
        # Caption Style
        if "IEEE Caption" not in [s.name for s in styles]:
            caption_style = styles.add_style("IEEE Caption", WD_STYLE_TYPE.PARAGRAPH)
            caption_style.font.name = self.FONT_TITLE
            caption_style.font.size = self.FONT_SIZE_CAPTION
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_before = Pt(6)
            caption_style.paragraph_format.space_after = Pt(12)
        
        # Keywords Style
        if "IEEE Keywords" not in [s.name for s in styles]:
            kw_style = styles.add_style("IEEE Keywords", WD_STYLE_TYPE.PARAGRAPH)
            kw_style.font.name = self.FONT_TITLE
            kw_style.font.size = self.FONT_SIZE_ABSTRACT
            kw_style.font.italic = True
            kw_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            kw_style.paragraph_format.space_after = Pt(12)
    
    def _set_two_column(self):
        """Add a section break and set two-column layout."""
        # Add section break for two columns
        new_section = self.doc.add_section()
        new_section.page_width = self.PAGE_WIDTH
        new_section.page_height = self.PAGE_HEIGHT
        new_section.top_margin = self.MARGIN_TOP
        new_section.bottom_margin = self.MARGIN_BOTTOM
        new_section.left_margin = self.MARGIN_LEFT
        new_section.right_margin = self.MARGIN_RIGHT
        
        # Set two columns using OOXML
        sectPr = new_section._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), str(int(self.COLUMN_GAP.twips)))
        sectPr.append(cols)
    
    def _add_title(self, title: str):
        """Add paper title."""
        p = self.doc.add_paragraph(title, style="IEEE Title")
    
    def _add_authors(self, authors: str):
        """Add author information."""
        p = self.doc.add_paragraph(authors, style="IEEE Author")
    
    def _add_abstract(self, abstract: str):
        """Add abstract section."""
        # Abstract label
        p = self.doc.add_paragraph(style="IEEE Abstract Label")
        run = p.add_run("Abstract—")
        run.bold = True
        run.italic = True
        run.add_text(abstract)
        p.style = "IEEE Abstract"
    
    def _add_keywords(self, keywords: List[str]):
        """Add keywords."""
        p = self.doc.add_paragraph(style="IEEE Keywords")
        run = p.add_run("Keywords—")
        run.bold = True
        run.italic = True
        p.add_run("; ".join(keywords))
    
    def _add_section_heading(self, heading: str, number: int):
        """Add a section heading with Roman numeral."""
        roman_numerals = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
        numeral = roman_numerals[number - 1] if number <= len(roman_numerals) else str(number)
        
        p = self.doc.add_paragraph(style="IEEE Heading")
        run = p.add_run(f"{numeral}. {heading.upper()}")
        run.bold = True
    
    def _add_body_text(self, content: str):
        """Add body text with proper formatting."""
        # Split content into paragraphs
        paragraphs = content.strip().split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Clean up the text
                para_text = ' '.join(para_text.split())
                p = self.doc.add_paragraph(para_text, style="IEEE Body")
    
    def _add_figure(self, image_data: bytes, caption: str):
        """Add a figure with caption."""
        self.figure_counter += 1
        
        # Add image
        try:
            image_stream = io.BytesIO(image_data)
            self.doc.add_picture(image_stream, width=Inches(3.0))
            
            # Center the last paragraph (which contains the image)
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            # If image fails, add placeholder text
            p = self.doc.add_paragraph("[Image could not be processed]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        p = self.doc.add_paragraph(style="IEEE Caption")
        run = p.add_run(f"Fig. {self.figure_counter}. ")
        run.bold = True
        p.add_run(caption)
    
    def _add_references(self, references: List[str]):
        """Add references section."""
        self._add_section_heading("References", 99)  # Will show as "REFERENCES"
        
        # Override the last heading to not show number
        last_para = self.doc.paragraphs[-1]
        last_para.clear()
        run = last_para.add_run("REFERENCES")
        run.bold = True
        
        for i, ref in enumerate(references, 1):
            p = self.doc.add_paragraph(style="IEEE Body")
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.add_run(f"[{i}] {ref}")
    
    def generate(
        self, 
        paper: StructuredPaper,
        images: Optional[List[bytes]] = None
    ) -> bytes:
        """
        Generate IEEE-formatted DOCX document.
        
        Args:
            paper: StructuredPaper object with all content
            images: Optional list of image bytes
            
        Returns:
            Bytes of the generated DOCX file
        """
        images = images or []
        
        # Set up document
        self._setup_document()
        
        # Add title and authors (single column)
        self._add_title(paper.title)
        self._add_authors(paper.authors)
        
        # Add abstract and keywords (still single column for now)
        self._add_abstract(paper.abstract)
        self._add_keywords(paper.keywords)
        
        # Switch to two-column layout
        self._set_two_column()
        
        # Build a map of figures by section
        figures_by_section = {}
        for fig in paper.figures:
            section_name = fig.placement.lower()
            if section_name not in figures_by_section:
                figures_by_section[section_name] = []
            figures_by_section[section_name].append(fig)
        
        # Add sections
        for i, section in enumerate(paper.sections, 1):
            self._add_section_heading(section.heading, i)
            self._add_body_text(section.content)
            
            # Add figures for this section
            section_key = section.heading.lower()
            if section_key in figures_by_section:
                for fig in figures_by_section[section_key]:
                    if fig.index < len(images):
                        self._add_figure(images[fig.index], fig.caption)
        
        # Add any remaining images not placed in sections
        placed_indices = {fig.index for fig in paper.figures}
        for i, img_data in enumerate(images):
            if i not in placed_indices:
                self.figure_counter += 1
                self._add_figure(img_data, f"Figure {self.figure_counter}")
        
        # Add references if present
        if paper.references:
            self._add_references(paper.references)
        
        # Save to bytes
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        
        return output.getvalue()


# Singleton instance
_generator: Optional[IEEEDocxGenerator] = None

def get_docx_generator() -> IEEEDocxGenerator:
    """Get or create the document generator singleton."""
    global _generator
    if _generator is None:
        _generator = IEEEDocxGenerator()
    return _generator
